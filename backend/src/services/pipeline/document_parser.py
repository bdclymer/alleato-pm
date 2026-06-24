"""
Stage 1 (Alternative) – Generic document parser.

Handles PDF and DOCX files (not meeting transcripts).
Extracts raw text, creates semantic segments via LLM,
and stores in document_metadata + meeting_segments
(reusing the same tables so Stage 2/3 work unmodified).

Usage:
    from .document_parser import run_document_parser
    run_document_parser(metadata_id="some-uuid")
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
from typing import Any, Dict, List, Optional

from ..supabase_helpers import (
    fetch_optional_row,
    get_rag_read_client,
    get_rag_write_client,
    get_supabase_client,
    update_ingestion_job_state,
)
from .models import MeetingSegment
from . import llm

logger = logging.getLogger(__name__)

# Max chars to send to LLM for segmentation (prevents token overflow)
MAX_LLM_CHARS = 30_000
MIN_EXTRACTED_CHARS = 50
DOC_SEGMENT_WINDOW_LINES = int(os.getenv("DOC_SEGMENT_WINDOW_LINES", "260"))
DOC_SEGMENT_WINDOW_OVERLAP = int(os.getenv("DOC_SEGMENT_WINDOW_OVERLAP", "40"))
DOC_SUMMARY_MAX_CHARS = int(os.getenv("DOC_SUMMARY_MAX_CHARS", "12000"))
DOC_SEGMENT_USE_LLM = (os.getenv("DOC_SEGMENT_USE_LLM", "true").strip().lower() not in {"0", "false", "no", "off"})


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: List[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
    return "\n\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip() or "1"
                paragraphs.append(f"{'#' * int(level)} {text}")
            else:
                paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        rows_text: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            paragraphs.append("\n".join(rows_text))

    return "\n\n".join(paragraphs)


def extract_text_from_plain(content: str) -> str:
    """Pass-through for plain text or markdown."""
    return content.strip()


def _sanitize_text(text: str) -> str:
    """Remove DB-breaking control chars while preserving formatting."""
    if not text:
        return ""
    return text.replace("\x00", "").replace("\u0000", "")


def detect_and_extract(
    file_bytes: Optional[bytes],
    content: Optional[str],
    file_name: Optional[str],
) -> str:
    """Detect document type and extract text."""
    ext = (file_name or "").lower().rsplit(".", 1)[-1] if file_name else ""

    if file_bytes:
        if ext == "pdf":
            return _sanitize_text(extract_text_from_pdf(file_bytes))
        elif ext in ("docx", "doc"):
            return _sanitize_text(extract_text_from_docx(file_bytes))
        else:
            # Try to decode as text
            try:
                return _sanitize_text(file_bytes.decode("utf-8").strip())
            except UnicodeDecodeError:
                # Fall back to PDF detection by magic bytes
                if file_bytes[:5] == b"%PDF-":
                    return _sanitize_text(extract_text_from_pdf(file_bytes))
                raise ValueError(f"Cannot extract text from file: {file_name}")

    if content:
        return _sanitize_text(extract_text_from_plain(content))

    raise ValueError("No file bytes or content provided")


# ---------------------------------------------------------------------------
# Document segmentation via LLM
# ---------------------------------------------------------------------------

def _fallback_window_segment(numbered_content: str, title: str) -> List[Dict[str, Any]]:
    """Return one deterministic segment for a content window."""
    indexes = [int(match) for match in re.findall(r"^\[(\d+)\]", numbered_content, re.MULTILINE)]
    if not indexes:
        return []
    return [
        {
            "title": title,
            "start_index": min(indexes),
            "end_index": max(indexes),
            "summary": f"Deterministic segment for {title}.",
            "decisions": [],
            "risks": [],
            "tasks": [],
        }
    ]


def _segment_document_window(
    numbered_content: str, title: str
) -> List[Dict[str, Any]]:
    """Segment one numbered content window."""
    if not DOC_SEGMENT_USE_LLM:
        return _fallback_window_segment(numbered_content, title)

    prompt = f"""Analyze this document and identify distinct semantic sections.

Document: {title}

Content (each line prefixed with [index]):
{numbered_content}

Return JSON array of sections. Each section should capture a coherent topic or content block.

Required format:
{{
  "segments": [
    {{
      "title": "Brief descriptive title for this section",
      "start_index": 0,
      "end_index": 15,
      "summary": "2-3 sentence summary of what this section covers",
      "decisions": ["Any decisions or conclusions stated in this section"],
      "risks": ["Any risks, issues, or concerns mentioned"],
      "tasks": ["Any action items, requirements, or next steps"]
    }}
  ]
}}

Guidelines:
- Sections should be 10-100 lines typically
- Every line must belong to exactly one section
- Capture natural content transitions (headings, topic changes)
- Include document header/intro and conclusion sections
- Extract any actionable items, decisions, or risks mentioned
- For specifications: each spec section is a segment
- For contracts: each clause/article is a segment
- For reports: each major section is a segment"""

    raw = llm._call_llm(prompt, json_mode=True)
    if not raw or not raw.strip():
        logger.warning(
            "[DocParser] LLM returned empty segmentation response for %s; using fallback segment",
            title,
        )
        return _fallback_window_segment(numbered_content, title)

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as exc:
        logger.warning(
            "[DocParser] Invalid segmentation JSON for %s; using fallback segment: %s",
            title,
            exc,
        )
        return _fallback_window_segment(numbered_content, title)
    segments = parsed.get("segments", [])
    if not isinstance(segments, list):
        logger.warning(
            "[DocParser] Segmentation response for %s did not include a segments array; using fallback segment",
            title,
        )
        return _fallback_window_segment(numbered_content, title)
    return segments


def _build_document_summary_excerpt(text: str) -> str:
    """Build representative excerpt from start/middle/end for long docs."""
    if len(text) <= DOC_SUMMARY_MAX_CHARS:
        return text

    sample = max(1200, DOC_SUMMARY_MAX_CHARS // 3)
    middle_start = max(0, (len(text) // 2) - (sample // 2))
    parts = [
        ("Start", text[:sample]),
        ("Middle", text[middle_start : middle_start + sample]),
        ("End", text[-sample:]),
    ]
    return "\n\n".join(f"[{label} Excerpt]\n{chunk}" for label, chunk in parts)


def _coerce_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _segment_document(text: str, title: str) -> List[Dict[str, Any]]:
    """Use LLM to semantically segment a document into logical sections."""
    lines = text.split("\n")
    if not lines:
        return []

    segments: List[Dict[str, Any]] = []
    start = 0
    while start < len(lines):
        end = min(len(lines), start + DOC_SEGMENT_WINDOW_LINES)
        window_lines = lines[start:end]
        if not window_lines:
            break

        numbered = "\n".join(
            f"[{start + i}] {line}" for i, line in enumerate(window_lines)
        )
        if len(numbered) > MAX_LLM_CHARS:
            numbered = numbered[:MAX_LLM_CHARS]

        window_title = f"{title} (lines {start}-{max(start, end - 1)})"
        window_segments = _segment_document_window(numbered, window_title)
        for seg in window_segments:
            seg_start = _coerce_int(seg.get("start_index"), start)
            seg_end = _coerce_int(seg.get("end_index"), max(start, end - 1))
            seg_start = max(start, min(end - 1, seg_start))
            seg_end = max(seg_start, min(end - 1, seg_end))
            segments.append(
                {
                    "title": seg.get("title", f"Section {seg_start}-{seg_end}"),
                    "start_index": seg_start,
                    "end_index": seg_end,
                    "summary": seg.get("summary", ""),
                    "decisions": seg.get("decisions") or [],
                    "risks": seg.get("risks") or [],
                    "tasks": seg.get("tasks") or [],
                }
            )

        if end >= len(lines):
            break
        start = max(0, end - DOC_SEGMENT_WINDOW_OVERLAP)

    deduped: List[Dict[str, Any]] = []
    seen_keys: set[tuple[Any, ...]] = set()
    for seg in sorted(
        segments,
        key=lambda s: (
            _coerce_int(s.get("start_index"), 0),
            _coerce_int(s.get("end_index"), 0),
        ),
    ):
        key = (
            _coerce_int(seg.get("start_index"), 0),
            _coerce_int(seg.get("end_index"), 0),
            str(seg.get("title", "")),
            str(seg.get("summary", "")),
        )
        if key in seen_keys:
            continue
        seen_keys.add(key)
        deduped.append(seg)
    return deduped


def _generate_document_summary(text: str, title: str) -> str:
    """Generate a summary of the document."""
    excerpt = _build_document_summary_excerpt(text)
    prompt = f"""Generate a comprehensive summary of this document.

Document: {title}

Content:
{excerpt}

Write a 2-4 paragraph summary covering:
1. Document purpose and type (spec, contract, report, etc.)
2. Key topics and sections
3. Important details, numbers, dates, or requirements
4. Action items or next steps if applicable

Be specific and include concrete details where present."""

    return llm._call_llm(prompt)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def run_document_parser(metadata_id: str) -> Dict[str, Any]:
    """
    Parse a generic document (PDF, DOCX, or text) and segment it.

    This is the alternative Stage 1 for non-meeting documents.
    It produces the same output (meeting_segments + document_metadata update)
    so Stage 2 (embedder) and Stage 3 (extractor) work without changes.

    Returns:
        dict with metadataId, segmentCount, summaryLength, extractedChars
    """
    client = get_supabase_client()

    # 1. Fetch metadata
    resp = (
        client.table("document_metadata")
        .select("id,title,type,category,source,source_system,project_id,date,captured_at,created_at,summary,overview,status,fireflies_id,participants,participants_array,storage_bucket,file_path,file_name,url,source_web_url,source_metadata,content,raw_text")
        .eq("id", metadata_id)
        .single()
        .execute()
    )
    metadata = resp.data
    if not metadata:
        raise ValueError(f"document_metadata not found: {metadata_id}")

    title = metadata.get("title") or "Untitled Document"
    category = metadata.get("category") or "document"
    logger.info("[DocParser] Processing: %s (%s) [%s]", title, metadata_id, category)

    # 2. Get text content — either from stored content or from file storage
    rag_metadata = fetch_optional_row(
        get_rag_read_client(),
        "rag_document_metadata",
        "content,raw_text",
        "id",
        metadata_id,
    )
    content = rag_metadata.get("content") or rag_metadata.get("raw_text") or metadata.get("content") or metadata.get("raw_text")
    file_bytes: Optional[bytes] = None

    # If no inline content, try to fetch from Supabase storage
    file_path = metadata.get("file_path")
    file_name = metadata.get("file_name") or metadata.get("title") or ""
    if not content and file_path:
        logger.info("[DocParser] Fetching file from storage: %s", file_path)
        bucket_name = metadata.get("storage_bucket") or "documents"
        download_resp = client.storage.from_(bucket_name).download(file_path)
        if download_resp:
            file_bytes = download_resp
        else:
            raise ValueError(f"Failed to download file: {file_path}")

    # Fallback: extract storage path from public URL (e.g. drawing PDFs where file_path is null)
    if not content and not file_bytes:
        import re as _re
        public_url = metadata.get("url") or metadata.get("source_web_url") or ""
        m = _re.search(r"/storage/v1/object/public/([^/]+)/(.+)$", public_url)
        if m:
            bucket_name, storage_path = m.group(1), m.group(2)
            logger.info("[DocParser] Downloading via URL-extracted path: %s / %s", bucket_name, storage_path)
            download_resp = client.storage.from_(bucket_name).download(storage_path)
            if download_resp:
                file_bytes = download_resp

    # 3. Extract text
    extracted_text = _sanitize_text(detect_and_extract(file_bytes, content, file_name))
    extracted_len = len(extracted_text.strip())
    extracted_is_too_short = extracted_len < MIN_EXTRACTED_CHARS

    logger.info("[DocParser] Extracted %d chars from %s", len(extracted_text), file_name or "content")

    # 4. Store raw text if not already stored
    if not content:
        get_rag_write_client().table("rag_document_metadata").upsert(
            {
                "id": metadata_id,
                "app_document_id": metadata_id,
                "raw_text": extracted_text[:500_000],
                "content": extracted_text[:500_000],
                "content_length": len(extracted_text[:500_000]),
                "parsing_status": "raw_text_extracted",
            }
        ).execute()

    # 5. Generate document summary
    if extracted_is_too_short:
        summary = (
            f"Minimal extract for '{title}'. "
            f"Parsed content was only {extracted_len} characters and may require OCR or a different source format."
        )
        logger.warning(
            "[DocParser] Extracted text too short (%d chars) for %s; using fallback summary/segment",
            extracted_len,
            metadata_id,
        )
    else:
        summary = _generate_document_summary(extracted_text, title)
        logger.info("[DocParser] Generated summary: %d chars", len(summary))

    # 6. Segment the document via LLM
    if extracted_is_too_short:
        segment_dicts = []
    else:
        segment_dicts = _segment_document(extracted_text, title)
        logger.info("[DocParser] Created %d segments", len(segment_dicts))

    # If LLM returned no segments, create a single fallback segment
    if not segment_dicts:
        line_count = max(1, len(extracted_text.split("\n")))
        fallback_summary = summary[:500]
        fallback_text = extracted_text.strip()
        if not fallback_text:
            fallback_text = f"Document: {title}\n\nNo extractable text was found in the uploaded file."
            extracted_text = fallback_text
        segment_dicts = [{
            "title": title,
            "start_index": 0,
            "end_index": line_count - 1,
            "summary": fallback_summary,
            "decisions": [],
            "risks": [],
            "tasks": [],
        }]

    # 7. Convert to MeetingSegment objects (reusing the same model)
    segments: List[MeetingSegment] = [
        MeetingSegment(
            segment_index=i,
            title=s.get("title", f"Section {i + 1}"),
            start_index=s.get("start_index", 0),
            end_index=s.get("end_index", 0),
            summary=s.get("summary", ""),
            decisions=s.get("decisions") or [],
            risks=s.get("risks") or [],
            tasks=s.get("tasks") or [],
        )
        for i, s in enumerate(segment_dicts)
    ]

    # 8. Reset prior parser output, then upsert segments to meeting_segments.
    # This keeps re-indexed local files from leaving stale sections behind when
    # the source document shrinks or its headings change.
    client.table("meeting_segments").delete().eq("metadata_id", metadata_id).execute()
    for seg in segments:
        client.table("meeting_segments").upsert(
            {
                "metadata_id": metadata_id,
                "segment_index": seg.segment_index,
                "title": seg.title,
                "start_index": seg.start_index,
                "end_index": seg.end_index,
                "summary": seg.summary,
                "decisions": seg.decisions,
                "risks": seg.risks,
                "tasks": seg.tasks,
            },
            on_conflict="metadata_id,segment_index",
        ).execute()

    # 9. Update document_metadata with summary and advance status
    update_data: Dict[str, Any] = {
        "overview": summary,
        "status": "segmented",
    }
    # Store content for the embedder to use if we extracted from a file
    if file_bytes and not content:
        get_rag_write_client().table("rag_document_metadata").upsert(
            {
                "id": metadata_id,
                "app_document_id": metadata_id,
                "content": extracted_text[:500_000],
                "raw_text": extracted_text[:500_000],
                "content_length": len(extracted_text[:500_000]),
                "parsing_status": "segmented",
                "overview": summary,
            }
        ).execute()

    client.table("document_metadata").update(update_data).eq("id", metadata_id).execute()

    # 10. Advance job stage
    update_ingestion_job_state(
        metadata_id,
        stage="segmented",
        error_message=None,
        client=client,
    )

    return {
        "metadataId": metadata_id,
        "segmentCount": len(segments),
        "summaryLength": len(summary),
        "extractedChars": len(extracted_text),
    }
